from pathlib import Path
import re
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

INPUT_PATH = Path('STEM_REPORT.md')
OUTPUT_PATH = Path('STEM_REPORT_FULL.docx')

# Read markdown source
text = INPUT_PATH.read_text(encoding='utf-8')
lines = text.splitlines()

doc = Document()
section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(13)
style.paragraph_format.line_spacing = 1.4
style.paragraph_format.space_after = Pt(6)

# Ensure heading styles use Times New Roman
for style_name in ['Heading 1', 'Heading 2', 'Heading 3']:
    if style_name in doc.styles:
        hstyle = doc.styles[style_name]
        hstyle.font.name = 'Times New Roman'
        if style_name == 'Heading 1':
            hstyle.font.size = Pt(14)
        else:
            hstyle.font.size = Pt(13)

# Title page
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('BÁO CÁO ĐỀ TÀI STEM')
run.font.name = 'Times New Roman'
run.font.size = Pt(18)
run.bold = True

doc.add_paragraph()
doc.add_paragraph()

title_text = 'STUDYHUB — XÂY DỰNG ỨNG DỤNG WEB HỌC TẬP THÔNG MINH TÍCH HỢP THUẬT TOÁN LẶP LẠI NGẮT QUÃNG SM-2 HỖ TRỢ ÔN THI VÀ QUẢN LÝ KIẾN THỨC CHO HỌC SINH THPT'
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(title_text)
run.font.name = 'Times New Roman'
run.font.size = Pt(16)
run.bold = True

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Học sinh thực hiện: ...........................................')
run.font.name = 'Times New Roman'
run.font.size = Pt(14)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Giáo viên hướng dẫn: ......................................')
run.font.name = 'Times New Roman'
run.font.size = Pt(14)

doc.add_page_break()

# Table of contents page
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('MỤC LỤC')
run.font.name = 'Times New Roman'
run.font.size = Pt(16)
run.bold = True

toc_paragraph = doc.add_paragraph()
field = OxmlElement('w:fldSimple')
field.set(qn('w:instr'), 'TOC \\o "1-3" \\h \\z \\u')
toc_paragraph._p.append(field)

doc.add_page_break()

# Helper functions
paragraph_buffer = []
current_table = []


def flush_paragraph():
    global paragraph_buffer
    if not paragraph_buffer:
        return
    paragraph_text = ' '.join(paragraph_buffer).strip()
    if paragraph_text:
        p = doc.add_paragraph(paragraph_text)
        p.style = doc.styles['Normal']
    paragraph_buffer = []


def flush_table():
    global current_table
    if not current_table:
        return
    rows = [row.strip().strip('|').split('|') for row in current_table]
    if rows:
        table = doc.add_table(rows=1, cols=len(rows[0]))
        table.style = 'Light List Accent 1'
        hdr_cells = table.rows[0].cells
        for idx, cell_text in enumerate(rows[0]):
            hdr_cells[idx].text = cell_text.strip()
            hdr_cells[idx].paragraphs[0].runs[0].font.name = 'Times New Roman'
            hdr_cells[idx].paragraphs[0].runs[0].font.size = Pt(12)
            hdr_cells[idx].paragraphs[0].runs[0].bold = True
        for row in rows[1:]:
            cells = table.add_row().cells
            for idx, cell_text in enumerate(row):
                cells[idx].text = cell_text.strip()
                cells[idx].paragraphs[0].runs[0].font.name = 'Times New Roman'
                cells[idx].paragraphs[0].runs[0].font.size = Pt(12)
    current_table = []


in_code_block = False
code_lines = []

for line in lines:
    stripped = line.strip()
    if stripped.startswith('```'):
        if in_code_block:
            if code_lines:
                p = doc.add_paragraph('\n'.join(code_lines))
                run = p.runs[0]
                run.font.name = 'Courier New'
                run.font.size = Pt(11)
            code_lines = []
        in_code_block = not in_code_block
        continue
    if in_code_block:
        code_lines.append(line)
        continue
    if stripped == '':
        flush_table()
        flush_paragraph()
        continue
    if line.startswith('|'):
        current_table.append(line)
        continue
    flush_table()
    if line.startswith('## '):
        flush_paragraph()
        doc.add_paragraph(line[3:].strip(), style='Heading 1')
        continue
    if line.startswith('### '):
        flush_paragraph()
        doc.add_paragraph(line[4:].strip(), style='Heading 2')
        continue
    if line.startswith('#### '):
        flush_paragraph()
        doc.add_paragraph(line[5:].strip(), style='Heading 3')
        continue
    if stripped.startswith('- ') or stripped.startswith('* '):
        flush_paragraph()
        p = doc.add_paragraph(stripped[2:].strip(), style='List Bullet')
        p.style.font.name = 'Times New Roman'
        p.style.font.size = Pt(13)
        continue
    if re.match(r'^\d+\.', stripped):
        flush_paragraph()
        p = doc.add_paragraph(stripped, style='List Number')
        p.style.font.name = 'Times New Roman'
        p.style.font.size = Pt(13)
        continue
    paragraph_buffer.append(line)

flush_table()
flush_paragraph()

# Footer page numbering
footer = doc.sections[0].footer
footer_para = footer.paragraphs[0]
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_para.clear()
run = footer_para.add_run('Trang ')
run.font.name = 'Times New Roman'
run.font.size = Pt(11)

fld_begin = OxmlElement('w:fldChar')
fld_begin.set(qn('w:fldCharType'), 'begin')
run._r.append(fld_begin)

instr = OxmlElement('w:instrText')
instr.set(qn('xml:space'), 'preserve')
instr.text = 'PAGE'
run._r.append(instr)

fld_separate = OxmlElement('w:fldChar')
fld_separate.set(qn('w:fldCharType'), 'separate')
run._r.append(fld_separate)

fld_end = OxmlElement('w:fldChar')
fld_end.set(qn('w:fldCharType'), 'end')
run._r.append(fld_end)

doc.save(OUTPUT_PATH)
print(f'Saved {OUTPUT_PATH.resolve()}')
